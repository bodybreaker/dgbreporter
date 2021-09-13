from argparse import ArgumentParser
import numpy as np
import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

kkma = Kkma()
okt = Okt()
stopwords = ['이상', '사항', '말', '말씀', 
             '이', '있', '하', '것', '들', '그', '되', '수', '이', '보', '않', '없', '나', 
             '사람', '주', '아니', '등', '같', '우리', '때', '년', '가', '한', '지', '대하', 
             '오', '말', '일', '그렇', '위하', '때문', '그것', '두', '말하', '알', '그러나', 
             '받', '못하', '일', '그런', '또', '문제', '더', '사회', '많', '그리고', '좋', '크', 
             '따르', '중', '나오', '가지', '씨', '시키', '만들', '지금', '생각하', '그러', '속', 
             '하나', '집', '살', '모르', '적', '월', '데', '자신', '안', '어떤', '내', '내', '경우', 
             '명', '생각', '시간', '그녀', '다시', '이런', '앞', '보이', '번', '나', '다른', '어떻', 
             '개', '전', '들', '사실', '이렇', '점', '싶', '말', '정도', '좀', '원', '잘', '통하', '소리', 
             '놓', '통해', '이후', '다음', '그래서', '왜', '응', '웅', '보자', '하자', '이다', '왜', '게', 
             '잖아','하자', '까', '일단', '우선', '에이', '비', '씨']

class SentenceTokenizer(object):
    def __init__(self):
        #self.url = url
        self.sentences = []
        #self.nouns = []
  
    def text2sentences(self, text):
        self.sentences = kss.split_sentences(text)
        for idx in range(0, len(self.sentences)):
            if len(self.sentences[idx]) <= 10:
                self.sentences[idx-1] += (' ' + self.sentences[idx])
                self.sentences[idx] = ''
        return self.sentences

    def get_nouns(self, sentences):
        nouns = []
        for sentence in sentences:
            if sentence != '':
                nouns.append(' '.join([noun for noun in okt.nouns(str(sentence)) 
                                       if noun not in stopwords and len(noun) > 1]))
        return nouns

class GraphMatrix(object):
    def __init__(self):
        self.tfidf = TfidfVectorizer()
        self.cnt_vec = CountVectorizer()
        self.graph_sentence = []
        
    def build_sent_graph(self, sentence):
        tfidf_mat = self.tfidf.fit_transform(sentence).toarray()
        self.graph_sentence = np.dot(tfidf_mat, tfidf_mat.T)
        return  self.graph_sentence
        
    def build_words_graph(self, sentence):
        cnt_vec_mat = normalize(self.cnt_vec.fit_transform(sentence).toarray().astype(float), axis=0)
        vocab = self.cnt_vec.vocabulary_
        return np.dot(cnt_vec_mat.T, cnt_vec_mat), {vocab[word] : word for word in vocab}
     
    def draw_graph(self):
        graph = nx.from_numpy_matrix(self.graph_sentence, create_using=nx.MultiDiGraph())
        pos = nx.circular_layout(graph)
        nx.draw_circular(graph)
        labels = {i : i for i in graph.nodes()}
        nx.draw_networkx_labels(graph, pos, labels, font_size=15)
        plt.show()

class Rank(object):
    def get_ranks(self, graph, d=0.85): # d = damping factor
        A = graph
        matrix_size = A.shape[0]
        for id in range(matrix_size):
            A[id, id] = 0 # diagonal 부분을 0으로 
            link_sum = np.sum(A[:,id]) # A[:, id] = A[:][id]
            if link_sum != 0:
                A[:, id] /= link_sum
            A[:, id] *= -d
            A[id, id] = 1
            
        B = (1-d) * np.ones((matrix_size, 1))
        ranks = np.linalg.solve(A, B) # 연립방정식 Ax = b
        return {idx: r[0] for idx, r in enumerate(ranks)}

class TextRank(object):
    def __init__(self, text):
        self.sent_tokenize = SentenceTokenizer()
        self.sentences = self.sent_tokenize.text2sentences(text)
        
        self.nouns = self.sent_tokenize.get_nouns(self.sentences)
                    
        self.graph_matrix = GraphMatrix()
        self.sent_graph = self.graph_matrix.build_sent_graph(self.nouns)
        self.words_graph, self.idx2word = self.graph_matrix.build_words_graph(self.nouns)
        
        self.rank = Rank()
        self.sent_rank_idx = self.rank.get_ranks(self.sent_graph)
        self.sorted_sent_rank_idx = sorted(self.sent_rank_idx, key=lambda k: self.sent_rank_idx[k], reverse=True)
        
        self.word_rank_idx =  self.rank.get_ranks(self.words_graph)
        self.sorted_word_rank_idx = sorted(self.word_rank_idx, key=lambda k: self.word_rank_idx[k], reverse=True)
        
        
        
    def summarize(self, sent_num=3):  
        summary = []
        index=[]
        for idx in self.sorted_sent_rank_idx[:sent_num]:
            index.append(idx)
        
        index.sort()
        for idx in index:
            summary.append(self.sentences[idx])
        
        return summary
    
#     def highlight_setences(self, sent_num=3, dark=0.8):
#         weights = [self.sent_rank_idx[idx] for idx in self.sent_rank_idx]
#         weights = (weights - min(weights))/(max(weights) - min(weights) + 1e-4)
#         html = ''
#         fmt = ' <span style="background-color: #{0:x}{0:x}ff">{1}</span>'
#         for idx in range(len(self.sentences)):
#             if idx in self.sorted_sent_rank_idx[:sent_num]:
#                 c = int(256*((1.-dark)*(1.-self.sent_rank_idx[idx])+dark))
#             else:
#                 c = int(256*((1.-dark)*(1.-0)+dark))    
#             html += fmt.format(c,self.sentences[idx])
#         display(HTML(html))
    
    def keywords(self, word_num=5):
        rank = Rank()
        rank_idx = rank.get_ranks(self.words_graph)
        sorted_rank_idx = sorted(rank_idx, key=lambda k: rank_idx[k], reverse=True)
        
        keywords = []
        index=[]
        for idx in sorted_rank_idx[:word_num]:
            index.append(idx)
            
        #index.sort()
        for idx in index:
            keywords.append(self.idx2word[idx])
        
        return keywords


def extract(extracted_text):
    if extracted_text:
        textrank = TextRank(extracted_text)
        _summarizedSentence = textrank.summarize(sent_num=7)
        _keywords = textrank.keywords(word_num=10)
        extracted_text = '녹취 부분:\n' + extracted_text
        summarizedSentence = '핵심문장(5)\n' + '\n'.join(_summarizedSentence)    
        keywords = '키워드(10):\n' + ', '.join(_keywords)
        
        templ = docx.Document('./word/minutes.docx')
        templ.tables[0].rows[5].cells[2].paragraphs[0].text += "%s\n\n%s\n\n%s"%(extracted_text,summarizedSentence,keywords)
        templ.save('./word/minutes.docx')
    
